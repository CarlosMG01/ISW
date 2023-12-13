import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from flask import Blueprint, render_template, request, redirect, url_for, session, flash, make_response
from flask import Flask, render_template, request, jsonify, send_file, abort
from uuid import uuid4
from .bbdd import DatabaseManager
from .bbdd import confirmar_correo_en_bd, enviar_correo_verificacion, obtener_correo_desde_token,login, enviar_correo_restablecer
#guardar_documento, obtener_textos
from fpdf import FPDF
import fitz # pip install pymupdf
from docx import Document
from pygoogletranslation import Translator
from datetime import datetime
import re
import pytesseract
import os
from werkzeug.utils import secure_filename
from io import BytesIO
from uuid import uuid4
from PIL import Image
import PyPDF2
import sys
from flask_socketio import emit, SocketIO




sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

socketio = SocketIO()
app = Flask(__name__)



docker = 0

auth_bp = Blueprint('auth', __name__)
home_bp = Blueprint('home', __name__)
chat_bp = Blueprint('chat', __name__)

resultado_global = ""

if docker == 1:
    db_manager = DatabaseManager('mysql-container', 'root', 'root', 'prueba')
else:
    db_manager = DatabaseManager('localhost', 'carlos', 'root', 'prueba')

def obtener_id_usuario_actual():
    correo = session.get('correo_usuario')
    if correo:
        return db_manager.obtener_id_usuario(correo)
    else:
        return None

@home_bp.route('/')
def index():
    return redirect(url_for('auth.restricted'))

@auth_bp.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('auth.restricted'))

@auth_bp.route('/registro', methods=['GET', 'POST'])
def registro():
    error = None
    success = None
    if request.method == 'POST':
        correo = request.form['correo']
        contrasena = request.form['contrasena']
        confirmar_contrasena = request.form['confirmar_contrasena']
        error, success = db_manager.register_user(correo, contrasena, confirmar_contrasena)  
        if success:            
            enviar_correo_verificacion(correo, contrasena) 
            success= 'Se ha enviado un correo de confirmación a tu dirección de correo electrónico.'
            session['correo_usuario'] = correo
            return render_template('inicio_sesion.html', error=error, success=success)
    return render_template('registro.html', error=error, success=success)

@auth_bp.route('/confirmar-correo/<token>')
def confirmar_correo(token):
    correo = obtener_correo_desde_token(token)
    contrasena = request.args.get('contrasena')

    if correo and contrasena :
        cursor = db_manager.cursor
        connection = db_manager.connection

        confirmar_correo_en_bd(correo,contrasena, cursor, connection)
        success = "Su correo se ha validado con éxito"
        return render_template('inicio_sesion.html', correo=correo, success=success)
    else:
        return render_template('inicio_sesion.html', error='Token inválido o expirado')

@auth_bp.route('/inicio_sesion', methods=['GET', 'POST'])
def inicio_sesion():
    error = None

    if request.method == 'POST':
        correo = request.form['correo']
        contrasena = request.form['contrasena']

        if correo and contrasena:
            success, error_message = login(correo, contrasena, db_manager.cursor)
            if success:
                session['correo_usuario'] = correo
                session['logged_in'] = True
                return redirect(url_for('auth.restricted'))
            else:
                error = error_message or "Credenciales incorrectas"
        else:
            error = "Correo y contraseña son obligatorios"

    return render_template('inicio_sesion.html', error=error)

@auth_bp.route('/cambio_contrasena', methods=['GET', 'POST'])
def cambio_contrasena():
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    error = None
    success = None
    if request.method == 'POST':
        contrasena_actual = request.form['contrasena_actual']
        nueva_contrasena = request.form['nueva_contrasena']
        confirmar_nueva_contrasena = request.form['confirmar_nueva_contrasena']
        if nueva_contrasena != confirmar_nueva_contrasena:
            error = "Las contraseñas nuevas no coinciden"
        else:
            correo = session.get('correo_usuario')
            error, success = db_manager.change_password(correo, contrasena_actual, nueva_contrasena)

    return render_template('cambio_contrasena.html', error=error, success=success)

@auth_bp.route('/restricted', methods=['GET', 'POST'])
def restricted():
    global resultado_global
    error = None
    texto_extraido = ''
    translated = None
    if request.method == 'POST': 
        if 'file' in request.files:
            file = request.files['file']
            _, file_extension = os.path.splitext(file.filename)
            # Comprobar imagenes
            if file_extension.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                texto_extraido = pytesseract.image_to_string(Image.open(file))
            # Comprobar PDFs
            elif file_extension.lower() == '.pdf':
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    texto_extraido += page.extract_text()
            # Comprobar DOCX
            elif file_extension.lower() == '.docx':
                texto_extraido = extract_text_from_docx(file)
            else:
                error = "Tipo de archivo no admitido."

            resultado_global = texto_extraido
            usuario_id = obtener_id_usuario_actual()
            if usuario_id is not None:
                titulo = "Documento"
                db_manager.guardar_documento(usuario_id, titulo, texto_extraido)
        elif 'dest_lang' in request.form:
            # Traducción
            translator = Translator()
            idioma = request.form.get('dest_lang')
            if idioma == "Traducir al Ingles": 
                dest_lang="en"
            elif idioma == "Frances":
                dest_lang = "fr"
            elif idioma =="Portugues":
                dest_lang = "pt"
            elif idioma == "Italiano":
                dest_lang = "it"
            else:
                error = "Idioma no reconocido"
            if resultado_global:
                    translated = translator.translate(resultado_global, src="es", dest=dest_lang).text
            else:
                error = "No hay texto para traducir"
        elif 'PDF' in request.form:
            if resultado_global:
                error, success= convertir_a_pdf(resultado_global)
                if success:
                    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                    filename = f'mitexto_{timestamp}.pdf'
                    return send_file(success, as_attachment=True, download_name=filename, mimetype='application/pdf')
            else:
                error = "No hay texto para generar un PDF"
        elif 'WORD' in request.form:
            if resultado_global:
                error, success= convertir_a_word(resultado_global)
                if success:
                    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                    filename = f'mitexto_{timestamp}.docx'
                    return send_file(success, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            else:
                error = "No hay texto para generar un Word"
    return render_template('restricted.html', resultado = resultado_global,translated_text=translated, error = error)

# PDF
def convertir_a_pdf(resultado_global):
    pdf_output = BytesIO()
    try:
        texto_limpio = re.sub(r'[^\x20-\x7E]', '', resultado_global)
        pdf = FPDF(orientation='P', unit='mm', format='A4')
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(190, 10, txt='PDF extraído de OCRTeam', ln=True, align='C')
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(190, 10, txt=texto_limpio, border=0, align='L')
        pdf_output.write(pdf.output(dest='S').encode('latin-1'))
        pdf_output.seek(0)
        success = pdf_output
    except Exception as e:
        error = f"Error al generar el PDF: {str(e)}"
        pdf_output.close()
        return error, None 
    return None, pdf_output

# WORD
def convertir_a_word(resultado_global):
    texto_limpio = re.sub(r'[^\x20-\x7E]', '', resultado_global)
    docx_output = BytesIO()
    try:
        document = Document()
        document.add_heading('Documento Word extraído de OCRTeam', 0)
        document.add_paragraph(texto_limpio)
        document.save(docx_output)
        docx_output.seek(0)
    except Exception as e:
        error = f"Error al generar ek Word: {str(e)}"
        docx_output.close()
        return error, None
    return None, docx_output

#WORD
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text


#Perfil
@auth_bp.route('/perfil-usuario', methods =['GET','POST'])
def perfil():
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    else:
        correo = session.get('correo_usuario')
        imagen_perfil = db_manager.obtener_imagen_perfil(correo)
        if  request.method =='POST':
            nueva_imagen_perfil = request.files['nueva_imagen']
            if nueva_imagen_perfil:
                # Verifica que la extensión del archivo sea válida (puedes personalizar esto)
                if nueva_imagen_perfil.filename != '' and nueva_imagen_perfil.filename.endswith(('.jpg', '.png', '.jpeg', '.gif', '.bmp')):
                    # Actualiza la imagen de perfil en la base de datos
                    result = db_manager.guardar_imagen_perfil(correo, nueva_imagen_perfil)
                    imagen_perfil = db_manager.obtener_imagen_perfil(correo)
                    if result:
                        return render_template('perfil-usuario.html', imagen_perfil=imagen_perfil, correo=correo, success=result)
                else:
                    error = "Formato de imagen no válido. Por favor, utiliza archivos de imagen (.jpg, .png, .jpeg, .gif, .bmp)."
        return render_template('perfil-usuario.html', imagen_perfil=imagen_perfil, correo=correo)


@auth_bp.route('/mistextos', methods=['GET'])
def mistextos():
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    else:
        correo = session.get('correo_usuario')
        usuario_id = obtener_id_usuario_actual()
        if usuario_id is not None:
            documentos = db_manager.obtener_documentos(usuario_id)
            return render_template('mistextos.html', correo=correo, documentos=documentos)
        else:
            return "Error al obtener ID del usuario."

@auth_bp.route('/generar-pdf/<int:id>', methods=['GET'])
def generar_pdf(id):
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    usuario_id = obtener_id_usuario_actual()
    if usuario_id is not None:
        documento = db_manager.obtener_documento_por_id(usuario_id, id)
        if documento is not None:
            contenido = documento[1]
            error, success= convertir_a_pdf(contenido)
            if success:
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                filename = f'mitexto_{timestamp}.pdf'
                return send_file(success, as_attachment=True, download_name=filename, mimetype='application/pdf')
        else:
            return "Documento no encontrado."
    else:
        return "Error al obtener ID del usuario."

@auth_bp.route('/generar-word/<int:id>', methods=['GET'])
def generar_word(id):
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    usuario_id = obtener_id_usuario_actual()

    if usuario_id is not None:
        documento = db_manager.obtener_documento_por_id(usuario_id, id)
        if documento is not None:
            contenido = documento[1]
            error, success = convertir_a_word(contenido)
            if success:
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                filename = f'mitexto{timestamp}.docx'
                return send_file(success, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            error = "Documento no encontrado."
            return error
    else:
        error = "Error al obtener ID del usuario."
        return error
    
@auth_bp.route('/borrar-texto/<int:id>', methods=['GET'])
def borrar_texto(id):
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    usuario_id = obtener_id_usuario_actual()

    if usuario_id is not None:

        documento = db_manager.obtener_documento_por_id(usuario_id, id)

        if documento is not None:
            db_manager.borrar_documento(id)
            flash('Texto borrado correctamente', 'success')
        else:
            flash('No se puede borrar el texto. El documento no pertenece al usuario actual.', 'error')
    else:
        flash('Error al obtener ID del usuario.', 'error')

    return redirect(url_for('auth.mistextos'))

@auth_bp.route('/editar-texto/<int:id>', methods=['GET', 'POST'])
def editar_texto(id):
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    usuario_id = obtener_id_usuario_actual()

    if usuario_id is not None:
        documento = db_manager.obtener_documento_por_id(usuario_id, id)

        if documento is not None:
            if request.method == 'POST':
                nuevo_contenido = request.form['nuevo_contenido']
                db_manager.actualizar_contenido_documento(id, nuevo_contenido)
                flash('Texto actualizado correctamente', 'success')
                return redirect(url_for('auth.mistextos'))

            return render_template('editar_texto.html', documento=documento)
        else:
            return "Documento no encontrado."
    else:
        return "Error al obtener ID del usuario."


@chat_bp.route('/chat')
def chat():
    messages = []   
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))

    correo = session.get('correo_usuario')
    
    return render_template('chat.html', correo=correo)

users = {}

@socketio.on("connect")
def handle_connect():
    print("Client connected!")

@socketio.on("user_join")
def handle_user_join(username):
    print(f"User {username} joined!")
    users[username] = request.sid

@socketio.on("new_message")
def handle_new_message(message):
    print(f"New message: {message}")
    username = None 
    for user in users:
        if users[user] == request.sid:
            username = user
    emit("chat", {"message": message, "username": username}, broadcast=True)


@auth_bp.route('/ayuda')
def ayuda():
    if not session.get('logged_in'):
        return redirect(url_for('auth.inicio_sesion'))
    return render_template('ayuda.html')


@auth_bp.route('/olvido_contrasena', methods=['GET', 'POST'])
def olvide_contrasena(): 
    error = None
    success = None
    if request.method == 'POST':
        correo = request.form['correo']
        if correo:
            error, success = db_manager.verificar_usuario_por_correo(correo)
            if success:
                enviar_correo_restablecer(correo)
    return render_template('olvido_contrasena.html', error = error, success=success)

@auth_bp.route('/restablecer_contrasena/<token>', methods=['GET', 'POST'])
def restablecer_contrasena(token):
    error = None
    success = None
    nueva_contrasena = None
    confirmar_contrasena = None

    correo = obtener_correo_desde_token(token)
    if request.method == 'POST':
        nueva_contrasena = request.form['nueva_contrasena']
        confirmar_contrasena = request.form['confirmar_nueva_contrasena']

        if nueva_contrasena != confirmar_contrasena:
           error = "Las contraseñas nuevas no coinciden"
        else:
            error, success = db_manager.restablecer_contrasena(correo, nueva_contrasena)
            if success:
                return render_template('inicio_sesion.html', success = success)
    return render_template('restablecer_contrasena.html', error=error, token=token)
